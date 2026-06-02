import os
import resend
from docx import Document
from datetime import datetime
import base64



resend.api_key = os.environ.get("RESEND_API_KEY")
def send_report_email(file_path, recipient_email):
    with open(file_path, "rb") as f:
        file_data = base64.b64encode(f.read()).decode("utf-8")

    response = resend.Emails.send({
        "from": f"AI Assistant <{os.environ.get('RESEND_EMAIL')}>",
        "to": recipient_email,
        "subject": "Your Business Report",
        "html": "<p>Attached is your latest business report.</p>",
        "attachments": [
            {
                "filename": "business_report.docx",
                "content": file_data
            }
        ]
    })

    return response



def generate_business_report(customers, sales):
    doc = Document()

    doc.add_heading("Business Performance Report", 0)

    doc.add_paragraph(f"Date: {datetime.now()}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Total Customers: {customers}")
    doc.add_paragraph(f"Total Sales: KES {sales}")

    doc.add_heading("Analysis", level=1)
    doc.add_paragraph(
        "This report summarizes the current business performance based on available data."
    )

    file_path = "business_report.docx"
    doc.save(file_path)

    return file_path